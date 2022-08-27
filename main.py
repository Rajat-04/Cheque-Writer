from num2words import num2words 
from PyPDF2 import PdfMerger

from tkinter import * 
from tkinter.ttk import *
from tkinter.filedialog import askopenfilename
from tkinter import Button
from tkinter import Checkbutton
from tkinter import Label
from tkinter import Entry
from tkinter import messagebox

import datetime
from datetime import date
from datetime import datetime

from openpyxl import load_workbook 

import customtkinter
import time
import os
import reportlab.rl_config 
import pyexcel as p
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib.pagesizes import inch
from reportlab.platypus import Table, TableStyle, Paragraph
from reportlab.lib.units import inch,cm
reportlab.rl_config.warnOnMissingFontGlyphs = 0

from pdf2docx import parse 
from docx import Document 


pdfmetrics.registerFont(TTFont('Calibri', 'Calibri.ttf'))
pdfmetrics.registerFont(TTFont('Calibri-BI', 'Calibriz.ttf'))
pdfmetrics.registerFont(TTFont('Calibri-B', 'Calibrib.ttf'))

lastdestroyed = set()
lastdestroyed2 = set()
lastdestroyed3 = set()
lastdestroyed4 = set()

desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')

def mainWindow():

	global ws
	global lastdestroyed
	lastdestroyed.clear()
	ws.title('Cheque Printing')
	ws.geometry('400x140')
	ws.eval('tk::PlaceWindow . center')
	ws.resizable(False, False)	



	def manauto():

		global lastdestroyed
		lastdestroyed2.clear()
		lastdestroyed3.clear()
		ws.title('Cheque Printing')
		ws.geometry('400x200')
		ws.eval('tk::PlaceWindow . center')
		ws.resizable(False, False)



		def clicked():
			backbutton.destroy()
			button3.destroy()
			button4.destroy()
			mainWindow()

		global backbutton		
		backbutton = Button(ws, command = clicked, text = '⮌', relief = 'groove')
		backbutton.pack(padx=10,pady=10, side = TOP, anchor = NW)


		def autoopenNormal():

			global lastdestroyed2,backbutton

			ws.title('Cheque Printing')
			ws.geometry('500x300')
			ws.eval('tk::PlaceWindow . center')
			ws.resizable(False, False)


			if 'button3' not in lastdestroyed2:
				lastdestroyed.update(('button3', 'button4', 'backbutton'))
				button3.destroy()
				button4.destroy()
				backbutton.destroy()




			def browseFiles():
				global file_path
				file_path = askopenfilename(initialdir = "/",title = "Select a File",filetypes = (("Excel Files","*.xlsx*"),("Excel Files","*.xls*"),("all files","*.*")))
				global checkprint
				checkprint = False
				if len(file_path)!=0 :
					if file_path[-5:]=='.xlsx' or file_path[-4:]=='.xls':
						if file_path[-4:]=='.xls':
							p.save_book_as(file_name=file_path,dest_file_name=file_path[:-4]+'.xlsx')
							file_path = file_path[:-4]+'.xlsx'						
						checkprint = True
						button6.configure(fg='white',bg='red' ,state = NORMAL)
						print(file_path)
						pb1 = Progressbar(ws, orient=HORIZONTAL, length=300, mode='determinate')
						pb1.place(relx=0.5, rely=0.85, anchor=CENTER)
						for i in range(5):
							ws.update_idletasks()
							pb1['value'] += 20
							time.sleep(0.05)
						pb1.destroy()
						label = Label(ws, text='File Uploaded Successfully! : ' + file_path, foreground='green')
						label.place(relx=0.5, rely=0.85, anchor=CENTER)
						ws.after(2000, lambda : label.destroy())

					else:
						button6.configure(fg='white', state = DISABLED)
						label = Label(ws, text='File Type Does Not Match Excel!', foreground='red')
						label.place(relx=0.5, rely=0.85, anchor=CENTER)
						ws.after(2000, lambda : label.destroy())



			def autonormalchequeprinting():

				global lastdestroyed

				def formatINR(number):
				    s, *d = str(number).partition(".")
				    r = ",".join([s[x-2:x] for x in range(-3, -len(s), -2)][::-1] + [s[-3:]])
				    return "".join([r] + d)


				def gen_Cheque(Name, Amount, Date):
				    fn = "pdf"
				    fn = fn + str(i)
				    fn = fn + ".pdf"
				    caw = num2words(Amount,lang='en_IN').upper() + " ONLY"
				    caw = caw.replace(',', '')
				    ca = "***"
				    ca = ca + str(formatINR(Amount)) + "/-"
				    Date = str(Date)[:10]
				    nowW = Date.split('-')
				    Date = date(int(nowW[0]),int(nowW[1]),int(nowW[2]))
				    day = str(Date.day)
				    if len(day) < 2:
				        day = '0 ' + str(day)
				    else:
				        day = day[0] + " " + day[1]
				    month = str(Date.month)
				    if len(month) < 2:
				        month = '0 ' + str(month)
				    else:
				        month = month[0] + " " + month[1]
				    year = str(Date.year)
				    year = year[0] + " " + year[1] + " " + year[2] + " " + year[3]
				    full_date = day + " / " + month + " / " + str(year)
				    canvas = Canvas(fn,pagesize = (8.66 * inch, 3.54 * inch))
				    canvas.setFont('Calibri', 11)
				    canvas.drawString(1.9 * inch,2.15 * inch, 'MENTOR FINMART PRIVATE LIMITED')
				    canvas.drawString(1.9 * inch,1.83 * inch, caw)
				    canvas.setFont('Calibri-BI', 11)
				    canvas.drawString(4.9 * inch,1.2 * inch, "A/C PAYEE")
				    '''canvas.setFont('Calibri-BI', 11)
				    canvas.drawString(4.9 * inch,1.12 * inch, "_________")'''
				    canvas.setFont('Calibri-B', 14)
				    canvas.drawString(6.7 * inch,1.52 * inch, ca)
				    canvas.setFont('Calibri', 12)
				    canvas.drawString(6.92 * inch,2.62 * inch, full_date)
				    if c1.get():
				    	canvas.drawImage('assets/logo.png',5.45 * inch,0.465 * inch,2.37 * cm,1.16 * cm, [0,0,0,0,0,0])
				    canvas.save()

				    
				wb = load_workbook(file_path)
				ws = wb.active
				Dates = ws['I']
				Date_list = [cell.value for cell in Dates[1:]]
				Amounts = ws['J']
				Amount_list = [cell.value for cell in Amounts[1:]]
				Names = ws['B']
				Name_list = [cell.value for cell in Names[1:]]


				for i in range(len(Name_list)):
				    gen_Cheque(Name_list[i], Amount_list[i], Date_list[i])


				pdfs = []

				for i in range(len(Name_list)):
					fn = "pdf"
					fn = fn + str(i)
					fn = fn + ".pdf"
					pdfs.append(fn)

				merger = PdfMerger()

				for pdf in pdfs:
				    merger.append(pdf)

				merger.write("result.pdf")
				merger.close()

				for pdf in pdfs:
					os.remove(pdf)
				
				if 'button5' not in lastdestroyed:
					lastdestroyed.update(('button5', 'button6', 'backbutton2'))
					backbutton2.destroy()
					button5.destroy()
					button6.destroy()
					ck1.destroy()

				pdftodoc()
				os.remove('result.pdf')
				os.remove(file_path)

				mainWindow()
				printfinal()



			def clicked():
				backbutton2.destroy()
				button5.destroy()
				button6.destroy()
				ck1.destroy()
				manauto()




			backbutton2 = backbutton = Button(ws, command = clicked, text = '⮌', relief = 'groove')
			backbutton2.pack(padx=10,pady=10, side = TOP, anchor = NW)

			button5 =  Button(master=ws, command=browseFiles, text='UPLOAD',bg = 'red',fg='white', width=14,height=1,font='Arial',activebackground='dodger blue',activeforeground='white',bd=0,pady=5,padx = 5)
			button5.place(relx=0.5, rely=0.5, anchor=CENTER)
			button5.pack(padx=10,pady=10)

			button6 =  Button(master=ws, command=autonormalchequeprinting, text='PRINT',bg = 'gray25',fg='white', width=14,height=1,font='Arial',activebackground='dodger blue',activeforeground='white',bd=0,pady=5,padx = 5,state= DISABLED)
			button6.place(relx=0.5, rely=0.5, anchor=CENTER)
			button6.pack(padx=10,pady=10)


			c1 = IntVar()
			ck1 = Checkbutton(ws, text = "Print Logo", 
	                      variable = c1,
	                      onvalue = 1,
	                      offvalue = 0,
	                      height = 2,
	                      width = 100)
			ck1.deselect()
			ck1.pack(padx=10,pady=10)



		def manualopenNormal():



			global lastdestroyed3,backbutton

			ws.title('Cheque Printing')
			#ws.geometry('1170x750')
			ws.state('zoomed')
			ws.resizable(False, False)
			




			if 'button3' not in lastdestroyed3:
				lastdestroyed.update(('button3', 'button4', 'backbutton'))
				backbutton.destroy()
				button3.destroy()
				button4.destroy()


			lo = Label(master=ws,text='', width=27, height=2, fg="#f0f0f0", bg= "#f0f0f0")
			ll = Label(master=ws,text='', width=27, height=2, fg="#f0f0f0", bg= "#f0f0f0")

			l1 = Label(master=ws,text='', width=14, height=2, fg="#f0f0f0", bg= "#f0f0f0")

			l2 = Label(master=ws,text='', width=14, height=2, fg="#f0f0f0", bg= "#f0f0f0")

			l3 = Label(master=ws,text='', width=14, height=2, fg="#f0f0f0", bg= "#f0f0f0")

			l4 = Label(master=ws,text='', width=14, height=2, fg="#f0f0f0", bg= "#f0f0f0")

			l5 = Label(master=ws,text='FULL NAME', width=14, height=2, fg="white", bg = "black", font='Arial')

			l6 = Label(master=ws,text='AMOUNT', width=14, height=2, fg="white", bg = "black", font='Arial')

			l7 = Label(master=ws,text='DATE', width=14, height=2, fg="white", bg = "black", font='Arial')

			l8 = Label(master=ws,text='', width=14, height=2, fg="#f0f0f0", bg= "#f0f0f0")			

			# grid method to arrange labels in respective
			# rows and columns as specified
			lo.grid(row = 0, column = 0, sticky = NS, pady = 2)
			l1.grid(row = 0, column = 1, sticky = NS, pady = 2)
			l2.grid(row = 0, column =3, sticky = NS, pady = 2)
			l3.grid(row = 0, column = 5, sticky = NS, pady = 2)
			l4.grid(row = 0, column = 7, sticky = NS, pady = 2)
			
			ll.grid(row = 1, column = 0, sticky = NS, pady = 2)
			l5.grid(row = 1, column = 1, sticky = NS, pady = 20)
			l6.grid(row = 1, column =3, sticky = NS, pady = 20)
			l7.grid(row = 1, column = 5, sticky = NS, pady = 20)
			l8.grid(row = 1, column = 7, sticky = NS, pady = 20)
			
			


			Nameeobj_list = []
			Amounteeobj_list = []
			Dateeobj_list = []
			logo_list = []
			vars_list = []



			def placeholder_name(event):
				if (event.widget.get())=='Enter Name':
					event.widget.delete(0, END)

			def placeholder_amount(event):
			    if (event.widget.get())=='Enter Amount':
			    	event.widget.delete(0, END)

			def placeholder_date(event):
			    if (event.widget.get())=='DD/MM/YYYY':
			    	event.widget.delete(0, END)



			# entry widgets, used to take entry from user
			name = Entry(master=ws, width=24,bd=2,fg='black', bg='white',font=('Arial',16),justify = CENTER, relief = 'groove')
			name.grid(row=0,column=1,ipady=11)
			name.insert(0,'Enter Name')
			Nameeobj_list.append(name)

			amount = Entry(master=ws, width=24,bd=2,fg='black', bg='white',font=('Arial',16),justify = CENTER, relief = 'groove')
			amount.grid(row=0,column=1,ipady=11)
			amount.insert(0,'Enter Amount')
			Amounteeobj_list.append(amount)

			date1 = Entry(master=ws, width=24,bd=2,fg='black', bg='white',font=('Arial',16),justify = CENTER, relief = 'groove')
			date1.grid(row=0,column=1,ipady=11)
			date1.insert(0,'DD/MM/YYYY')
			Dateeobj_list.append(date1)


			Nameeobj_list[0].grid(row = 2, column = 1, pady = 20, padx = 23)
			Nameeobj_list[0].bind("<Button-1>", placeholder_name)
			Nameeobj_list[0].widget = 'n0'
			Amounteeobj_list[0].grid(row = 2, column = 3, pady = 20, padx = 23)
			Amounteeobj_list[0].bind("<Button-1>", placeholder_amount)
			Amounteeobj_list[0].widget = 'a0'
			Dateeobj_list[0].grid(row = 2, column = 5, pady = 20, padx = 23)
			Dateeobj_list[0].bind("<Button-1>", placeholder_date)
			Dateeobj_list[0].widget = 'd0'

			c1 = IntVar()
			ck1 = Checkbutton(ws, text = "Print Logo", 
	                      variable = c1,
	                      onvalue = 1,
	                      offvalue = 0,
	                      height = 1,
	                      width = 10,
	                      font=10)
			ck1.deselect()
			logo_list.append(ck1)
			vars_list.append(c1)
			logo_list[0].grid(row = 2, column = 7, pady = 20)
			

			



			global count
			count = 1


			def addRecord():

				count = len(Nameeobj_list)
				print(Nameeobj_list[0].get())

				if count >= 5:

					AddButton.configure(state=DISABLED,fg='gray25',bg='gray25')
					DeleteButton.configure(state=NORMAL,fg='white',bg='red')

				else:
					DeleteButton.configure(state=NORMAL,fg='white',bg='red')

					name = Entry(master=ws, width=24,bd=2,fg='black', bg='white',font=('Arial',16),justify = CENTER, relief = 'groove')
					name.grid(row=0,column=1,ipady=11)
					name.insert(0,'Enter Name')
					Nameeobj_list.append(name)

					amount = Entry(master=ws, width=24,bd=2,fg='black', bg='white',font=('Arial',16),justify = CENTER, relief = 'groove')
					amount.grid(row=0,column=1,ipady=11)
					amount.insert(0,'Enter Amount')
					Amounteeobj_list.append(amount)

					date2 = Entry(master=ws, width=24,bd=2,fg='black', bg='white',font=('Arial',16),justify = CENTER, relief = 'groove')
					date2.grid(row=0,column=1,ipady=11)
					date2.insert(0,'DD/MM/YYYY')
					Dateeobj_list.append(date2)

					c1 = IntVar()
					ck1 = Checkbutton(ws, text = "Print Logo", 
			                      variable = c1,
			                      onvalue = 1,
			                      offvalue = 0,
			                      height = 1,
			                      width = 10,
			                      font=10)
					ck1.deselect()
					logo_list.append(ck1)
					vars_list.append(c1)
					


					Nameeobj_list[count].grid(row = count + 2, column = 1, pady = 20, padx = 23)
					Nameeobj_list[count].bind("<Button-1>", placeholder_name)
					Nameeobj_list[count].widget = 'n' + str(count)
					Amounteeobj_list[count].grid(row = count  + 2, column = 3, pady = 20, padx = 23)
					Amounteeobj_list[count].bind("<Button-1>", placeholder_amount)
					Amounteeobj_list[count].widget = 'a' + str(count)
					Dateeobj_list[count].grid(row = count + 2, column = 5, pady = 20, padx = 23)
					Dateeobj_list[count].bind("<Button-1>", placeholder_date)
					Dateeobj_list[count].widget = 'd' + str(count)
					logo_list[count].grid(row = count + 2, column = 7, pady = 20)

					count = count + 1

					if count >= 5:
						AddButton.configure(state=DISABLED,fg='gray25',bg='gray25')
					

			def deleteRecord():
				count = len(Nameeobj_list)

				if count == 1:
					AddButton.configure(state=NORMAL,fg='white',bg='red')
					DeleteButton.configure(state=DISABLED,fg='gray25',bg='gray25')

				else:
					AddButton.configure(state=NORMAL,fg='white',bg='red')
					Nameeobj_list[-1].destroy()
					Amounteeobj_list[-1].destroy()
					Dateeobj_list[-1].destroy()
					logo_list[-1].destroy()

					Nameeobj_list.pop(-1)
					Amounteeobj_list.pop(-1)
					Dateeobj_list.pop(-1)
					logo_list.pop(-1)

					count = count - 1

					if count == 1:
						DeleteButton.configure(state=DISABLED,fg='gray25',bg='gray25')


			def setEntries():

				a, b = validate()

				if a:

					global Namee_list, Amounte_list, Datee_list
					Namee_list = []
					Amounte_list = []
					Datee_list = []

					for i in range(len(Nameeobj_list)):
						Namee_list.append(Nameeobj_list[i].get())

					for i in range(len(Amounteeobj_list)):
						Amounte_list.append(Amounteeobj_list[i].get())

					for i in range(len(Dateeobj_list)):
						if '/' in Dateeobj_list[i].get():
							day, month, year = Dateeobj_list[i].get().split('/')
						elif '.' in Dateeobj_list[i].get():
							day, month, year = Dateeobj_list[i].get().split('.')
						if day[0] == '0':
							day = day[1] 
						if month[0] == '0':
							month = month[1] 
						Datee_list.append(date(int(year), int(month), int(day)))

					res = messagebox.askquestion('Confirm Print', 'Do you want to continue to print?')
					
					if res == 'yes':

						manualnormalchequeprinting()
						time.sleep(1)
						
						if 'l1' not in lastdestroyed:
							l1.destroy()
							l2.destroy()
							l3.destroy()
							l4.destroy()
							l5.destroy()
							l6.destroy()
							l7.destroy()
							l8.destroy()
							for i in range(len(Nameeobj_list)):
								Nameeobj_list[i].destroy()
								Amounteeobj_list[i].destroy()
								Dateeobj_list[i].destroy()
								logo_list[i].destroy()
							backbutton3.destroy()
							AddButton.destroy()
							DeleteButton.destroy()
							printButton.destroy()
							lo.destroy()
							ll.destroy()

						mainWindow()

					elif res == 'no':
						pass
					else:
						messagebox.showwarning('error', 'Something went wrong!')
				
				else:
					messagebox.showwarning(b[0], b[1])


			


			def clicked():
				l1.destroy()
				l2.destroy()
				l3.destroy()
				l4.destroy()
				l5.destroy()
				l6.destroy()
				l7.destroy()
				l8.destroy()
				for i in range(len(Nameeobj_list)):
					Nameeobj_list[i].destroy()
					Amounteeobj_list[i].destroy()
					Dateeobj_list[i].destroy()
					logo_list[i].destroy()
				backbutton3.destroy()
				AddButton.destroy()
				DeleteButton.destroy()
				printButton.destroy()
				lo.destroy()
				ll.destroy()
				manauto()




			def manualnormalchequeprinting():

				def formatINR(number):
				    s, *d = str(number).partition(".")
				    r = ",".join([s[x-2:x] for x in range(-3, -len(s), -2)][::-1] + [s[-3:]])
				    return "".join([r] + d)


				def gen_Cheque(Name, Amount, Date, logo):
				    fn = "pdf"
				    fn = fn + str(i)
				    fn = fn + ".pdf"
				    caw = num2words(Amount,lang='en_IN').upper() + " ONLY"
				    caw = caw.replace(',', '')
				    ca = "***"
				    ca = ca + str(formatINR(Amount)) + "/-"


				    day = str(Date.day)
				    if len(day) < 2:
				        day = '0 ' + str(day)
				    else:
				        day = day[0] + " " + day[1]
				    month = str(Date.month)
				    if len(month) < 2:
				        month = '0 ' + str(month)
				    else:
				        month = month[0] + " " + month[1]
				    year = str(Date.year)
				    year = year[0] + " " + year[1] + " " + year[2] + " " + year[3]
				    full_date = day + " / " + month + " / " + str(year)
				    canvas = Canvas(fn,pagesize = (8.66 * inch, 3.54 * inch))
				    canvas.setFont('Calibri', 11)
				    canvas.drawString(1.9 * inch,2.15 * inch, Name.upper())
				    canvas.drawString(1.9 * inch,1.83 * inch, caw)
				    '''canvas.drawString(4.9 * inch,1.2 * inch, "_________")'''
				    canvas.setFont('Calibri-BI', 11)
				    canvas.drawString(4.9 * inch,1.2 * inch, "A/C PAYEE")
				    canvas.setFont('Calibri-B', 14)
				    canvas.drawString(6.7 * inch,1.52 * inch, ca)
				    canvas.setFont('Calibri', 12)
				    canvas.drawString(6.92 * inch,2.62 * inch, full_date)
				    if vars_list[i].get():
				    	canvas.drawImage('assets/logo.png',5.45 * inch,0.465 * inch,2.37 * cm,1.16 * cm, [0,0,0,0,0,0])
				    canvas.save()


				for i in range(len(Namee_list)):
				    gen_Cheque(Namee_list[i], Amounte_list[i], Datee_list[i], vars_list[i])


				pdfs = []

				for i in range(len(Namee_list)):
					fn = "pdf"
					fn = fn + str(i)
					fn = fn + ".pdf"
					pdfs.append(fn)


				
				merger = PdfMerger()

				for pdf in pdfs:
				    merger.append(pdf)

				merger.write("result.pdf")
				merger.close()

				for pdf in pdfs:
					os.remove(pdf)

				pdftodoc()
				os.remove('result.pdf')
				printfinal()
			


			def validate():
				flag1 , flag2, flag3 = True, True, True
				s = ['Field Empty', 'Fill all the fields before print!']
				for i in range(len(Nameeobj_list)):
					if Nameeobj_list[i].get() == '' or Nameeobj_list[i].get() == 'Enter Name':
						flag1 = False
						break

				for i in range(len(Amounteeobj_list)):
					if Amounteeobj_list[i].get() == '' or Amounteeobj_list[i].get() == 'Enter Amount':
						flag2 = False
						break
					elif not Amounteeobj_list[i].get().isdigit():
						flag2 = False
						s[0] = 'Invalid Amount Type'
						s[1] = 'Account Number must only contain digits'
						break
				
				for i in range(len(Dateeobj_list)):
					if Dateeobj_list[i].get() == '' or Dateeobj_list[i].get() == 'DD/MM/YYYY':
						flag3 = False
						break
					elif 1:
						try:
							if '/' in Dateeobj_list[i].get():
								datetime.strptime(Dateeobj_list[i].get(), '%d/%m/%Y')
							elif '.' in Dateeobj_list[i].get():
								datetime.strptime(Dateeobj_list[i].get(), '%d.%m.%Y')
						except ValueError:
						  s = ['Invalid Date', 'Enter Correct Date Format!']
						  flag3=False
						  break
		
				ans =  flag1 and flag2 and flag3
				return ans, s
					

			backbutton3 = Button(ws, command = clicked, text = '⮌', relief = 'groove')
			backbutton3.place(x=20,y=20, anchor = NW)


			AddButton =  Button(master=ws,command=addRecord,width=16,height=2,text='ADD RECORD', font="Arial",fg='white',bg='red',activebackground='dodger blue',activeforeground='white',bd=0,pady=5,padx = 5)
			AddButton.grid(row = 7, column = 1, pady = 20, padx = 20)

			DeleteButton =  Button(master=ws,command=deleteRecord,width=16,height=2,text='DELETE RECORD', font="Arial",fg='gray25',bg='gray25',state=DISABLED,activebackground='dodger blue',activeforeground='white',bd=0,pady=5,padx = 5)
			DeleteButton.grid(row = 7, column = 5, pady = 20, padx = 20)
			
			printButton =  Button(master=ws,command=setEntries,width=16,height=2,text='PRINT', font="Arial",fg='white',bg='red',state=NORMAL,activebackground='dodger blue',activeforeground='white',bd=0,pady=5,padx = 5)
			printButton.grid(row = 7, column = 3, pady = 20, padx = 20)





		button3 =  Button(master=ws,command=autoopenNormal, text='AUTOMATIC',bg = 'red',fg='white', width=14,height=1,font='Arial',activebackground='dodger blue',activeforeground='white',bd=0,pady=5,padx = 5)
		button3.place(relx=0.5, rely=0.5, anchor=CENTER)
		button3.pack(padx=10,pady=10)

		button4 =  Button(master=ws,command=manualopenNormal,text='MANUAL',bg = 'red',fg='white', width=14,height=1,font='Arial',activebackground='dodger blue',activeforeground='white',bd=0,pady = 5, padx = 5)
		button4.place(relx=0.5, rely=0.5, anchor=CENTER)
		button4.pack(padx=10,pady=10)


		if 'button1' not in lastdestroyed:
			lastdestroyed.clear()
			lastdestroyed.update(('button1', 'button2'))
			button1.destroy()
			button2.destroy()




		




	def openRTGS():

		global lastdestroyed
		ws.title('Cheque Printing')
		ws.state('zoomed')
		ws.resizable(False, False)




		if 'button1' not in lastdestroyed:
			lastdestroyed.clear()
			lastdestroyed.update(('button1', 'button2'))
			button1.destroy()
			button2.destroy()




		l1 = Label(master=ws,text='', width=14, height=2, fg="#f0f0f0", bg= "#f0f0f0")

		l2 = Label(master=ws,text='', width=14, height=2, fg="#f0f0f0", bg= "#f0f0f0")
		
		l3 = Label(master=ws,text='', width=14, height=2, fg="#f0f0f0", bg= "#f0f0f0")

		l4 = Label(master=ws,text='', width=14, height=2, fg="#f0f0f0", bg= "#f0f0f0")

		l5 = Label(master=ws,text='', width=14, height=2, fg="#f0f0f0", bg= "#f0f0f0")

		l6 = Label(master=ws,text='A/c Holder Name', width=16, height=2, fg="white", bg = "black", font='Arial')

		l7 = Label(master=ws,text='A/c Number', width=16, height=2, fg="white", bg = "black", font='Arial')
		
		l8 = Label(master=ws,text='Bank', width=16, height=2, fg="white", bg = "black", font='Arial')

		l9 = Label(master=ws,text='IFSC Code', width=16, height=2, fg="white", bg = "black", font='Arial')

		l10 = Label(master=ws,text='Amount', width=16, height=2, fg="white", bg = "black", font='Arial')
		




		# grid method to arrange labels in respective
		# rows and columns as specified

		l1.grid(row = 0, column = 1, sticky = NS, pady = 2)
		l2.grid(row = 0, column =3, sticky = NS, pady = 2)
		l3.grid(row = 0, column = 5, sticky = NS, pady = 2)
		l4.grid(row = 0, column =7, sticky = NS, pady = 2)
		l5.grid(row = 0, column = 9, sticky = NS, pady = 2)

		l6.grid(row = 1, column = 1, sticky = NS, pady = 20)
		l7.grid(row = 1, column =3, sticky = NS, pady = 20)
		l8.grid(row = 1, column = 5, sticky = NS, pady = 20)
		l9.grid(row = 1, column =7, sticky = NS, pady = 20)
		l10.grid(row = 1, column = 9, sticky = NS, pady = 20)
		



		global Nameeobj_list,Amounteeobj_list,Bankobj_list,IFSCobj_list,numberobj_list
		Nameeobj_list = []
		Amounteeobj_list = []
		Bankobj_list = []
		IFSCobj_list = []
		numberobj_list = []



		# entry widgets, used to take entry from user

		for i in range(1):

			name = Entry(master=ws, width=20,bd=2,fg='black', bg='white',font=('Arial',16),justify = CENTER, relief = 'groove')
			name.grid(row=0,column=1,ipady=11)

			Nameeobj_list.append(name)

			number = Entry(master=ws, width=20,bd=2,fg='black', bg='white',font=('Arial',16),justify = CENTER, relief = 'groove')
			number.grid(row=0,column=1,ipady=11)

			numberobj_list.append(number)

			bank = Entry(master=ws, width=29,bd=2,fg='black', bg='white',font=('Arial',16),justify = CENTER, relief = 'groove')
			bank.grid(row=0,column=1,ipady=11)

			Bankobj_list.append(bank)

			ifsc = Entry(master=ws, width=20,bd=2,fg='black', bg='white',font=('Arial',16),justify = CENTER, relief = 'groove')
			ifsc.grid(row=0,column=1,ipady=11)
			IFSCobj_list.append(ifsc)

			amount = Entry(master=ws, width=20,bd=2,fg='black', bg='white',font=('Arial',16),justify = CENTER, relief = 'groove')
			amount.grid(row=0,column=1,ipady=11)

			Amounteeobj_list.append(amount)

			Nameeobj_list[i].grid(row = i + 2, column = 1, pady = 20, padx = 20)

			numberobj_list[i].grid(row = i + 2, column = 3, pady = 20, padx = 20)

			Bankobj_list[i].grid(row = i + 2, column = 5, pady = 20, padx = 20)
			
			IFSCobj_list[i].grid(row = i + 2, column = 7, pady = 20, padx = 20)
			
			Amounteeobj_list[i].grid(row = i + 2, column = 9, pady = 20, padx = 20)



		def validate():

			flag = False
			
			s = ['Field Empty', 'Fill all the fields before print!']
			
			if not Amounteeobj_list[0].get()=='':
				if not Amounteeobj_list[0].get().isdigit():
					s[0] = 'Invalid Amount Type'
					s[1] = 'Amount must only contain digits'
				else:
					flag = True

			elif not IFSCobj_list[0].get()=='':
				if len(IFSCobj_list[0].get())!=11:
					s[0] = 'Invalid IFSC'
					s[1] = 'Enter Valid IFSC'
				else:
					flag = True

			elif not numberobj_list[0].get()=='':
				if not numberobj_list[0].get().isdigit():
					s[0] = 'Invalid A/c Number'
					s[1] = 'Account Number must only contain digits'
				else:
					flag = True

			return flag, s





		def setEntries():

			a, b = validate()

			if a:

				global Namee_list, Amounte_list, Bank_list, IFSC_list, ACnum_list
				printClicked = True
				Namee_list = []
				Amounte_list = []
				Bank_list = []
				IFSC_list = []
				ACnum_list = []

				for i in range(len(Nameeobj_list)):
					Namee_list.append(Nameeobj_list[i].get())
					ACnum_list.append(numberobj_list[i].get())
					Bank_list.append(Bankobj_list[i].get())
					IFSC_list.append(IFSCobj_list[i].get())
					Amounte_list.append(Amounteeobj_list[i].get())


				res = messagebox.askquestion('Confirm Print', 'Do you want to continue to print?')
				if res == 'yes':
					rtgsfrontcheque()
				elif res == 'no':
					pass
				else:
					messagebox.showwarning('error', 'Something went wrong!')

			else:
				messagebox.showwarning(b[0], b[1])





		def rtgsfrontcheque():

			global lastdestroyed

			totalamount = 0

			for i in range(len(Amounteeobj_list)):
				ac = Amounteeobj_list[i].get()
				counts = ac.count(',')
				if counts > 0:
					ac = list(ac)
					while counts :
					    ac.remove(',')
					    counts -= 1
					ac = ''.join(ac)
				if len(ac)>0:
					totalamount += int(ac)


			def formatINR(number):
			    s, *d = str(number).partition(".")
			    r = ",".join([s[x-2:x] for x in range(-3, -len(s), -2)][::-1] + [s[-3:]])
			    return "".join([r] + d)

			def gen_Cheque(totalamount):
			    
			    caw = num2words(totalamount,lang='en_IN').upper() + " ONLY"
			    caw = caw.replace(',', '')
			    ca = "***"
			    ca = ca + str(formatINR(totalamount)) + "/-"

			    d1 = customDate.get()
			    #d1 = datelist[7:]

			    d1 = d1.split('/')
			    d1 = ''.join(d1)
			    d1 = list(d1)
			    
			    full_date = d1[0]  + ' ' + d1[1] + " / " + d1[2]  + ' ' + d1[3] + " / " + d1[4]  + ' ' + d1[5] + ' ' + d1[6]  + ' ' + d1[7]  
			    canvas = Canvas('Front.pdf',pagesize = (8.66 * inch, 3.54 * inch))
			    canvas.setFont('Calibri', 11)
			    canvas.drawString(1.9 * inch,2.15 * inch, 'YOURSELF FOR RTGS/NEFT')
			    canvas.drawString(1.9 * inch,1.83 * inch, caw)
			    '''canvas.drawString(4.9 * inch,1.2 * inch, "_________")'''
			    canvas.setFont('Calibri-BI', 11)
			    canvas.drawString(4.9 * inch,1.2 * inch, "A/C PAYEE")
			    canvas.setFont('Calibri-B', 14)
			    canvas.drawString(6.7 * inch,1.52 * inch, ca)
			    canvas.setFont('Calibri', 12)
			    canvas.drawString(6.92 * inch,2.62 * inch, full_date)
			    if c1.get():
			    	canvas.drawImage('assets/logo.png',5.45 * inch,0.465 * inch,2.37 * cm,1.16 * cm, [0,0,0,0,0,0])
			    canvas.save()

			gen_Cheque(totalamount)

			def rtgsbackchequeprinting():

				def formatINR(number):
				    s, *d = str(number).partition(".")
				    r = ",".join([s[x-2:x] for x in range(-3, -len(s), -2)][::-1] + [s[-3:]])
				    return "".join([r] + d)


				def gen_Cheque(Name, Amount, Banks, IFSC, ACCOUNTNUMBER):
					
					styles = getSampleStyleSheet()
					style = styles["BodyText"]

					canv = Canvas("Back.pdf", pagesize=(8.66 * inch, 3.54 * inch))

					data = ['A/c HOLDER NAME', 'A/c NUMBER', 'BANK NAME', 'IFSC CODE', 'AMOUNT']
				    
					totalamount = 0

					for i in range(len(Name)):
						ac = Amount[i].strip()
						counts = ac.count(',')
						if counts > 0:
							ac = list(ac)
							while counts :
							    ac.remove(',')
							    counts -= 1
							ac = ''.join(ac)
						if len(ac)>0:
							totalamount += int(ac)

						acc = str(formatINR(ac)) + "/-" if len(ac)>0 else ''
					
					totalamount = 'Rs: ' + str(formatINR(totalamount)) + "/-"

					canv.setFont('Calibri-B', 14)
					canv.drawString(110,200, data[0] + " : " + Name[i].upper().strip())
					canv.setFont('Calibri-B', 14)
					canv.drawString(110,175, data[1] + " : " + ACCOUNTNUMBER[i].strip())
					canv.setFont('Calibri-B', 14)
					canv.drawString(110,150, data[2] + " : " + Banks[i].upper().strip())
					canv.setFont('Calibri-B', 14)
					canv.drawString(110,125, data[3] + " : " + IFSC[i].upper().strip())
					canv.setFont('Calibri', 14)
					canv.drawString(420,80, "For Mentor Finmart Pvt Ltd")
					canv.drawString(420,20, "Authorised Signature")
					canv.save()

				
				gen_Cheque(Namee_list, Amounte_list, Bank_list, IFSC_list, ACnum_list)
				time.sleep(2)
				pdfs = ['Front.pdf', 'Back.pdf']

				merger = PdfMerger()

				for pdf in pdfs:
				    merger.append(pdf)

				merger.write("result.pdf")
				merger.close()

				for pdf in pdfs:
					os.remove(pdf)
			
				
				if 'l1' not in lastdestroyed:
					l1.destroy()
					l2.destroy()
					l3.destroy()
					l4.destroy()
					l5.destroy()
					l6.destroy()
					l7.destroy()
					l8.destroy()
					l9.destroy()
					l10.destroy()
					for i in range(len(Nameeobj_list)):
						Nameeobj_list[i].destroy()
						numberobj_list[i].destroy()
						Bankobj_list[i].destroy()
						IFSCobj_list[i].destroy()
						Amounteeobj_list[i].destroy()		
					backbuttonrtgs.destroy()
					printfrontbackButton.destroy()
					ck1.destroy()
					customDate.destroy()




				mainWindow()

			
			rtgsbackchequeprinting()

			pdftodoc()
			os.remove('result.pdf')
			printfinal()






		def clicked():
			l1.destroy()
			l2.destroy()
			l3.destroy()
			l4.destroy()
			l5.destroy()
			l6.destroy()
			l7.destroy()
			l8.destroy()
			l9.destroy()
			l10.destroy()
			ck1.destroy()
			customDate.destroy()
			for i in range(len(Nameeobj_list)):
				Nameeobj_list[i].destroy()
				numberobj_list[i].destroy()
				Bankobj_list[i].destroy()
				IFSCobj_list[i].destroy()
				Amounteeobj_list[i].destroy()		
			backbuttonrtgs.destroy()
			printfrontbackButton.destroy()
			mainWindow()




		backbuttonrtgs = Button(ws, command = clicked, text = '⮌',relief='groove')
		backbuttonrtgs.place(x=15,y=15, anchor = NW)


		printfrontbackButton =  Button(master=ws,command=setEntries,width=16,height=2,text='Generate', font="Arial",fg='white',bg='red',state=NORMAL,activebackground='dodger blue',activeforeground='white',bd=0,pady=5,padx = 5)
		printfrontbackButton.grid(row = 8, column = 5, pady = 20, padx = 20)

		c1 = IntVar()
		ck1 = Checkbutton(ws, text = "Print Logo", 
                      variable = c1,
                      onvalue = 1,
                      offvalue = 0,
                      height = 1,
                      width = 10,
                      font=10)
		ck1.deselect()
		ck1.grid(row = 8, column = 7)


		currdate = date.today()
		currday = str(currdate.day)
		if currdate.day<10:
			currday = '0' + currday
		currmonth = str(currdate.month)
		if currdate.month<10:
			currmonth = '0' + currmonth
		curryear = str(currdate.year)
		cust = currday + '/' + currmonth + '/' + curryear

		customDate = customtkinter.CTkEntry(master=ws,
                               placeholder_text="Date : DD/MM/YYYY",
                               placeholder_text_color='black',
                               width=220,
                               height=25,
                               border_width=2,
                               corner_radius=10,
                               text_font=('Arial', 16),
                               text_color=('black','black'),
                               fg_color=('white','white'),
                               justify=CENTER)

		customDate.insert(0,cust)
		customDate.grid(row=8,column=3,ipady=11)





	button1 =  Button(master=ws, command=manauto,text='CHEQUE',bg = 'red',fg='white', width=10,height=1,font='Arial',activebackground='dodger blue',activeforeground='white',bd=0,pady=5,padx = 5)
	button1.place(relx=0.5, rely=0.5, anchor=CENTER)
	button1.pack(padx=10,pady=10)
	button2 =  Button(master=ws, command=openRTGS,text='RTGS',bg = 'red',fg='white', width=10,height=1,font='Arial',activebackground='dodger blue',activeforeground='white',bd=0,pady=5,padx = 5)
	button2.place(relx=0.5, rely=0.5, anchor=CENTER)
	button2.pack(padx=10,pady=10)
	


def printfinal():
	os.startfile('result.docx','open')


def printpreview():
	pass


def stay_on_top():
   ws.lift()
   ws.after(10, stay_on_top)



def pdftodoc():
	pdf_file = 'result.pdf'
	docx_file = 'result.docx'
	parse(pdf_file, docx_file)
	document = Document(docx_file)
	style = document.styles['Normal']
	font = style.font
	font.name = 'Calibri'
	for paragraph in document.paragraphs:
	    if 'A/C PAYEE' == paragraph.text:
	        paragraph.text = ''
	        paragraph.style = document.styles['Normal']
	        run = paragraph.add_run('A/C PAYEE')
	        run.underline = True
	        run.bold = True
	        run.italic = True

	document.save(docx_file)


	
if __name__ == '__main__':
	ws = Tk()
	ws.iconbitmap('assets/cheque.ico')
	mainWindow()
	stay_on_top()
	ws.mainloop()